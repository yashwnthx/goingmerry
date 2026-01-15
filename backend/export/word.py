"""Word Export"""
from datetime import datetime, timezone
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class WordExporter:
    HEADING_STYLES = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    
    def export(self, doc_model: dict) -> bytes:
        if doc_model.get("type") != "word":
            raise ValueError(f"Cannot export '{doc_model.get('type')}' as Word")
        
        doc = Document()
        self._set_properties(doc, doc_model)
        self._add_title(doc, doc_model.get("title", "Untitled"))
        
        for section in doc_model.get("sections", []):
            self._add_section(doc, section)
        
        self._add_sources(doc, doc_model.get("meta", {}).get("sources", []))
        return self._to_bytes(doc)
    
    def _set_properties(self, doc: Document, doc_model: dict) -> None:
        props = doc.core_properties
        props.author = "Merry"
        props.last_modified_by = "Merry"
        props.title = doc_model.get("title", "Untitled")
        
        meta = doc_model.get("meta", {})
        if meta.get("created_at"):
            try:
                props.created = datetime.fromisoformat(meta["created_at"].replace("Z", "+00:00"))
            except (ValueError, TypeError):
                props.created = datetime(2024, 1, 1, tzinfo=timezone.utc)
    
    def _add_title(self, doc: Document, title: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(24)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.space_after = Pt(24)
    
    def _add_section(self, doc: Document, section: dict, depth: int = 0) -> None:
        level = section.get("level", 1)
        doc.add_heading(section.get("heading", ""), level=level)
        
        if content := section.get("content"):
            para = doc.add_paragraph(content)
            if section.get("verification_status") == "needs_verification":
                run = para.add_run(" [Needs Verification]")
                run.italic = True
                run.font.size = Pt(9)
        
        for child in section.get("children", []):
            self._add_section(doc, child, depth + 1)
    
    def _add_sources(self, doc: Document, sources: list) -> None:
        if not sources:
            return
        doc.add_page_break()
        doc.add_heading("Sources", level=1)
        for i, source in enumerate(sources, 1):
            url = source.get("url", "Unknown")
            title = source.get("title", "")
            doc.add_paragraph(f"{i}. {title}: {url}" if title else f"{i}. {url}")
    
    def _to_bytes(self, doc: Document) -> bytes:
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()


def export_word_document(doc_model: dict) -> bytes:
    return WordExporter().export(doc_model)
